import base64
import io
import os
import re as _re_module
import zipfile
from datetime import datetime
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP

from django.conf import settings
from django.db import transaction
from django.http import HttpResponse
from django.db.models import Count, Sum, Value
from django.db.models.functions import Coalesce, TruncMonth
from django.shortcuts import get_object_or_404
from django.utils import timezone
from docx import Document as DocxDocument
from docx.table import _Row as _TableRow
from openpyxl import load_workbook
from rest_framework import status
from rest_framework.generics import ListAPIView, ListCreateAPIView, RetrieveUpdateDestroyAPIView
from rest_framework.parsers import FormParser, MultiPartParser
from rest_framework.response import Response
from rest_framework.views import APIView

from .models import (
    DocumentCalculation,
    DocumentCalculationCategory,
    NormativeCoefficient,
    OrganizationSettings,
)
from .serializers import (
    DocumentCalculationCreateSerializer,
    DocumentCalculationCategorySerializer,
    DocumentCalculationSerializer,
    HealthCheckSerializer,
    NormativeCoefficientSerializer,
    OrganizationSettingsSerializer,
)


class HealthCheckAPIView(APIView):
    authentication_classes = []
    permission_classes = []

    def get(self, request):
        payload = {
            "message": "app_main API ishlayapti",
            "service": "shnq-backend",
            "status": "ok",
        }
        serializer = HealthCheckSerializer(payload)
        return Response(serializer.data, status=status.HTTP_200_OK)


class NormativeCoefficientListAPIView(ListAPIView):
    authentication_classes = []
    permission_classes = []
    serializer_class = NormativeCoefficientSerializer
    queryset = NormativeCoefficient.objects.filter(is_active=True).order_by("normative_type")


class DocumentCalculationCategoryListAPIView(ListAPIView):
    authentication_classes = []
    permission_classes = []
    serializer_class = DocumentCalculationCategorySerializer
    queryset = DocumentCalculationCategory.objects.all().order_by("name", "id")


class DocumentCalculationListCreateAPIView(ListCreateAPIView):
    authentication_classes = []
    permission_classes = []
    queryset = DocumentCalculation.objects.all().order_by("-id")

    def get_serializer_class(self):
        if self.request.method == "POST":
            return DocumentCalculationCreateSerializer
        return DocumentCalculationSerializer

    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        instance = serializer.save()
        response_data = DocumentCalculationSerializer(instance).data
        return Response(response_data, status=status.HTTP_201_CREATED)


class DocumentCalculationRetrieveUpdateDestroyAPIView(RetrieveUpdateDestroyAPIView):
    authentication_classes = []
    permission_classes = []
    queryset = DocumentCalculation.objects.all()

    def get_serializer_class(self):
        if self.request.method in ("PUT", "PATCH"):
            return DocumentCalculationCreateSerializer
        return DocumentCalculationSerializer

    def update(self, request, *args, **kwargs):
        partial = kwargs.pop("partial", False)
        instance = self.get_object()
        serializer = self.get_serializer(instance, data=request.data, partial=partial)
        serializer.is_valid(raise_exception=True)
        instance = serializer.save()
        response_data = DocumentCalculationSerializer(instance).data
        return Response(response_data, status=status.HTTP_200_OK)


def _normalize_header(value) -> str:
    if value is None:
        return ""
    return str(value).strip().lower().replace(" ", "_")


def _to_text(value, default: str = "") -> str:
    if value is None:
        return default
    return str(value).strip()


def _to_int(value, default: int = 0) -> int:
    if value in (None, ""):
        return default
    try:
        return int(float(str(value).replace(",", ".").strip()))
    except (TypeError, ValueError):
        return default


def _to_decimal(value, default: Decimal = Decimal("0.00")) -> Decimal:
    if value in (None, ""):
        return default
    if isinstance(value, Decimal):
        return value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    normalized = str(value).replace("\u00a0", "").replace(" ", "").replace(",", ".").strip()
    try:
        return Decimal(normalized).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    except (InvalidOperation, ValueError):
        return default


def _q2(value: Decimal) -> Decimal:
    return value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)


class DocumentCalculationXlsxImportAPIView(APIView):
    authentication_classes = []
    permission_classes = []
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request):
        file_obj = request.FILES.get("file")
        if not file_obj:
            return Response({"detail": "XLSX fayl yuborilmadi."}, status=status.HTTP_400_BAD_REQUEST)
        if not file_obj.name.lower().endswith(".xlsx"):
            return Response({"detail": "Faqat .xlsx formatdagi fayl qo'llab-quvvatlanadi."}, status=status.HTTP_400_BAD_REQUEST)

        try:
            workbook = load_workbook(file_obj, data_only=True)
            worksheet = workbook.active
        except Exception:
            return Response({"detail": "XLSX faylni o'qib bo'lmadi."}, status=status.HTTP_400_BAD_REQUEST)

        rows_iter = worksheet.iter_rows(values_only=True)
        header_row = next(rows_iter, None)
        if not header_row:
            return Response({"detail": "XLSX fayl bo'sh."}, status=status.HTTP_400_BAD_REQUEST)

        headers = [_normalize_header(cell) for cell in header_row]
        header_index = {header: idx for idx, header in enumerate(headers) if header}
        required_headers = [
            "category",
            "name",
            "development_deadline",
            "executor_organization",
            "notes",
            "total_pages",
            "normative_type",
            "document_category",
            "complexity_level",
            "completed_amount",
        ]
        missing_headers = [header for header in required_headers if header not in header_index]
        if missing_headers:
            return Response(
                {"detail": "XLSX ustunlari to'liq emas.", "missing_headers": missing_headers},
                status=status.HTTP_400_BAD_REQUEST,
            )

        final_total_header = next(
            (
                key
                for key in ("final_total_amount", "final_total", "total_amount", "umumiy_narxi")
                if key in header_index
            ),
            None,
        )
        normative_values = {choice[0] for choice in NormativeCoefficient.NormativeType.choices}
        document_category_values = {choice[0] for choice in DocumentCalculation.DocumentCategory.choices}
        complexity_values = {choice[0] for choice in DocumentCalculation.ComplexityLevel.choices}

        created_count = 0
        errors = []

        with transaction.atomic():
            for row_number, row in enumerate(rows_iter, start=2):
                if row is None:
                    continue
                if all(value in (None, "") for value in row):
                    continue

                get_value = lambda key: row[header_index[key]] if header_index[key] < len(row) else None

                name = _to_text(get_value("name"))
                if not name:
                    continue

                category_name = _to_text(get_value("category"))
                development_deadline = _to_text(get_value("development_deadline"))
                executor_organization = _to_text(get_value("executor_organization"))
                notes = _to_text(get_value("notes"))
                total_pages = max(_to_int(get_value("total_pages"), default=0), 0)
                complexity_level = _to_text(get_value("complexity_level"))
                if complexity_level.endswith(".0"):
                    complexity_level = complexity_level[:-2]

                raw_normative = _to_text(get_value("normative_type"))
                raw_document_category = _to_text(get_value("document_category"))

                # Fayldagi ustunlar joylashuvi almashib qolsa avtomatik tuzatamiz.
                if raw_normative in document_category_values and raw_document_category in normative_values:
                    raw_normative, raw_document_category = raw_document_category, raw_normative

                if raw_normative not in normative_values:
                    errors.append(
                        {
                            "row": row_number,
                            "field": "normative_type",
                            "message": f"Noto'g'ri qiymat: {raw_normative}",
                        }
                    )
                    continue
                if raw_document_category not in document_category_values:
                    errors.append(
                        {
                            "row": row_number,
                            "field": "document_category",
                            "message": f"Noto'g'ri qiymat: {raw_document_category}",
                        }
                    )
                    continue
                if complexity_level not in complexity_values:
                    errors.append(
                        {
                            "row": row_number,
                            "field": "complexity_level",
                            "message": f"Noto'g'ri qiymat: {complexity_level}",
                        }
                    )
                    continue

                completed_amount = _to_decimal(get_value("completed_amount"))
                final_total_from_xlsx = (
                    _to_decimal(row[header_index[final_total_header]], default=Decimal("0.00"))
                    if final_total_header and header_index[final_total_header] < len(row)
                    else None
                )
                category_obj = None
                if category_name:
                    category_obj, _ = DocumentCalculationCategory.objects.get_or_create(name=category_name)

                instance = DocumentCalculation(
                    calculation_category=category_obj,
                    name=name,
                    total_pages=total_pages,
                    normative_type=raw_normative,
                    document_category=raw_document_category,
                    complexity_level=complexity_level,
                    is_research_required=False,
                    development_deadline=development_deadline,
                    executor_organization=executor_organization,
                    notes=notes,
                    completed_amount=completed_amount,
                )
                instance.apply_normative_coefficients()

                if final_total_from_xlsx is not None:
                    instance.final_total_amount = final_total_from_xlsx
                else:
                    instance.recalculate_final_total_amount()

                # Shart: planned = final_total - completed, completed=0 bo'lsa planned=final_total bo'ladi.
                instance.planned_amount = (instance.final_total_amount - completed_amount).quantize(
                    Decimal("0.01"), rounding=ROUND_HALF_UP
                )
                instance.save()
                created_count += 1

            if errors:
                transaction.set_rollback(True)
                return Response(
                    {
                        "detail": "XLSX importda xatoliklar bor.",
                        "created_count": 0,
                        "error_count": len(errors),
                        "errors": errors,
                    },
                    status=status.HTTP_400_BAD_REQUEST,
                )

        return Response(
            {
                "detail": "XLSX ma'lumotlari muvaffaqiyatli import qilindi.",
                "created_count": created_count,
                "error_count": 0,
                "errors": [],
            },
            status=status.HTTP_201_CREATED,
        )


class DocumentCalculationReportTableAPIView(APIView):
    authentication_classes = []
    permission_classes = []

    def get(self, request):
        queryset = (
            DocumentCalculation.objects.select_related("calculation_category")
            .all()
            .order_by("calculation_category__name", "id")
        )

        sections_by_name: dict[str, dict] = {}
        sections: list[dict] = []
        total_amount_all = Decimal("0.00")
        completed_amount_all = Decimal("0.00")
        planned_amount_all = Decimal("0.00")
        next_year_amount_all = Decimal("0.00")

        for item in queryset:
            category_name = (
                item.calculation_category.name
                if item.calculation_category and item.calculation_category.name
                else "Kategoriyasiz"
            )
            if category_name not in sections_by_name:
                section_obj = {
                    "category": category_name,
                    "rows": [],
                    "totals": {
                        "total_amount": Decimal("0.00"),
                        "completed_amount": Decimal("0.00"),
                        "planned_amount": Decimal("0.00"),
                        "next_year_amount": Decimal("0.00"),
                    },
                }
                sections_by_name[category_name] = section_obj
                sections.append(section_obj)

            if item.current_year_percent:
                planned_amount = _q2(item.final_total_amount * item.current_year_percent / Decimal("100"))
                next_year_amount = _q2(
                    item.final_total_amount * max(Decimal("0"), Decimal("100") - item.current_year_percent) / Decimal("100")
                )
            else:
                planned_amount = _q2(item.final_total_amount - item.completed_amount)
                next_year_amount = Decimal("0.00")

            section = sections_by_name[category_name]
            order = len(section["rows"]) + 1
            section["rows"].append(
                {
                    "order": order,
                    "id": item.id,
                    "name": item.name,
                    "total_amount": _q2(item.final_total_amount),
                    "completed_amount": _q2(item.completed_amount),
                    "planned_amount": planned_amount,
                    "next_year_amount": next_year_amount,
                    "development_deadline": item.development_deadline,
                    "executor_organization": item.executor_organization,
                    "notes": item.notes,
                }
            )

            section["totals"]["total_amount"] += _q2(item.final_total_amount)
            section["totals"]["completed_amount"] += _q2(item.completed_amount)
            section["totals"]["planned_amount"] += planned_amount
            section["totals"]["next_year_amount"] += next_year_amount

            total_amount_all += _q2(item.final_total_amount)
            completed_amount_all += _q2(item.completed_amount)
            planned_amount_all += planned_amount
            next_year_amount_all += next_year_amount

        for section in sections:
            section["totals"]["total_amount"] = str(_q2(section["totals"]["total_amount"]))
            section["totals"]["completed_amount"] = str(_q2(section["totals"]["completed_amount"]))
            section["totals"]["planned_amount"] = str(_q2(section["totals"]["planned_amount"]))
            section["totals"]["next_year_amount"] = str(_q2(section["totals"]["next_year_amount"]))
            for row in section["rows"]:
                row["total_amount"] = str(_q2(row["total_amount"]))
                row["completed_amount"] = str(_q2(row["completed_amount"]))
                row["planned_amount"] = str(_q2(row["planned_amount"]))
                row["next_year_amount"] = str(_q2(row["next_year_amount"]))

        total_limit = _to_decimal(os.getenv("REPORT_TOTAL_LIMIT"), default=Decimal("0.00"))
        unallocated_limit = _q2(total_limit - planned_amount_all) if total_limit > planned_amount_all else Decimal("0.00")
        grand_total = _q2(planned_amount_all + unallocated_limit)

        payload = {
            "sections": sections,
            "summary": {
                "total_amount": str(_q2(total_amount_all)),
                "completed_amount": str(_q2(completed_amount_all)),
                "planned_amount": str(_q2(planned_amount_all)),
                "next_year_amount": str(_q2(next_year_amount_all)),
                "unallocated_limit": str(unallocated_limit),
                "grand_total": str(grand_total),
            },
        }
        return Response(payload, status=status.HTTP_200_OK)


class DashboardStatsAPIView(APIView):
    authentication_classes = []
    permission_classes = []

    def get(self, request):
        queryset = DocumentCalculation.objects.all()
        now = timezone.localtime()
        month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)

        type_choices = dict(NormativeCoefficient.NormativeType.choices)
        type_keys = list(type_choices.keys())

        by_type_rows = (
            queryset.values("normative_type")
            .annotate(
                count=Count("id"),
                total_amount=Coalesce(Sum("final_total_amount"), Value(Decimal("0.00"))),
            )
            .order_by("normative_type")
        )
        by_type_map = {row["normative_type"]: row for row in by_type_rows}

        this_month_rows = (
            queryset.filter(created_at__gte=month_start)
            .values("normative_type")
            .annotate(count=Count("id"))
            .order_by("normative_type")
        )
        this_month_map = {row["normative_type"]: row["count"] for row in this_month_rows}

        type_stats = []
        for key in type_keys:
            data = by_type_map.get(key, {})
            type_stats.append(
                {
                    "normative_type": key,
                    "label": type_choices[key],
                    "count": data.get("count", 0),
                    "this_month_count": this_month_map.get(key, 0),
                    "total_amount": data.get("total_amount", Decimal("0.00")),
                }
            )

        # Oxirgi 6 oy dinamikasi (jami hujjatlar soni)
        def shift_month(base: datetime, months_back: int) -> datetime:
            year = base.year
            month = base.month - months_back
            while month <= 0:
                month += 12
                year -= 1
            return base.replace(year=year, month=month, day=1)

        start_month = shift_month(month_start, 5)
        monthly_rows = (
            queryset.filter(created_at__gte=start_month)
            .annotate(month=TruncMonth("created_at"))
            .values("month")
            .annotate(count=Count("id"))
            .order_by("month")
        )
        monthly_count_map = {
            row["month"].strftime("%Y-%m"): row["count"] for row in monthly_rows if row["month"] is not None
        }

        trend = []
        month_labels = ["Yan", "Fev", "Mar", "Apr", "May", "Iyun", "Iyul", "Avg", "Sen", "Okt", "Noy", "Dek"]
        for offset in range(5, -1, -1):
            m = shift_month(month_start, offset)
            key = m.strftime("%Y-%m")
            trend.append(
                {
                    "key": key,
                    "label": month_labels[m.month - 1],
                    "count": monthly_count_map.get(key, 0),
                }
            )

        latest_rows = list(
            queryset.order_by("-created_at").values(
                "id",
                "name",
                "normative_type",
                "document_category",
                "final_total_amount",
                "created_at",
            )[:5]
        )

        totals = queryset.aggregate(
            total_documents=Count("id"),
            total_amount=Coalesce(Sum("final_total_amount"), Value(Decimal("0.00"))),
        )

        payload = {
            "totals": totals,
            "types": type_stats,
            "trend_last_6_months": trend,
            "latest_calculations": latest_rows,
        }
        return Response(payload, status=status.HTTP_200_OK)


import copy
import re

BOLD_PLACEHOLDERS = {
    "shnq_name",
    "shartnoma_number",
    "institute_director",
    "deputy_minister",
    "economics_head",
    "total_pages",
    "final_total_amount",
    "final_total_amount_words",
    "amount_2026",
    "amount_2027",
    "mhb_amount",
    "mhi_amount",
    "executor_organization",
    "development_deadline",
    "created_at",
    "notes",
    "normative_type",
}

# ---------------------------------------------------------------------------
# Raqamni o'zbek tilida so'zga o'girish (lotin)
# ---------------------------------------------------------------------------
_ONES = [
    "", "bir", "ikki", "uch", "to'rt", "besh",
    "olti", "yetti", "sakkiz", "to'qqiz",
]
_TENS = [
    "", "o'n", "yigirma", "o'ttiz", "qirq", "ellik",
    "oltmish", "yetmish", "sakson", "to'qson",
]


def _three_digits_to_words(n: int) -> str:
    """0–999 oralig'idagi sonni so'zga o'giradi."""
    if n == 0:
        return ""
    parts = []
    if n >= 100:
        h = n // 100
        parts.append(("" if h == 1 else _ONES[h] + " ") + "yuz")
        n %= 100
    if n >= 10:
        parts.append(_TENS[n // 10])
        n %= 10
    if n:
        parts.append(_ONES[n])
    return " ".join(parts)


def _fmt_money(value) -> str:
    """Summani mingliklar probel bilan ajratib, 2 xona kasr bilan formatlaydi.

    Misol: 511644672.00 → "511 644 672.00",  2369128.5 → "2 369 128.50"
    """
    try:
        q = Decimal(str(value)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    except (InvalidOperation, TypeError, ValueError):
        return str(value)
    return f"{q:,.2f}".replace(",", " ")  # '511 644 672.00'


def number_to_uz_words(amount) -> str:
    """
    Decimal/float/int sonni o'zbek tilida so'zga o'giradi.
    Misol: 276578765.50 → "Ikki yuz yetmish olti million besh yuz yetmish sakkiz ming
                           yetti yuz oltmish besh so'm ellik tiyin"
    """
    from decimal import Decimal as D
    amount = D(str(amount)).quantize(D("0.01"))
    integer_part = int(amount)
    tiyin_part = int(round((amount - integer_part) * 100))

    if integer_part == 0 and tiyin_part == 0:
        return "nol so'm"

    chunks = [
        (1_000_000_000_000, "trillion"),
        (1_000_000_000, "milliard"),
        (1_000_000, "million"),
        (1_000, "ming"),
    ]

    parts = []
    remaining = integer_part
    for divisor, name in chunks:
        if remaining >= divisor:
            q = remaining // divisor
            remaining %= divisor
            # "bir ming" emas, faqat "ming"
            if q == 1 and name == "ming":
                parts.append("ming")
            else:
                parts.append(_three_digits_to_words(q) + " " + name)

    if remaining:
        parts.append(_three_digits_to_words(remaining))

    result = " ".join(p.strip() for p in parts if p.strip())
    # Birinchi harfni katta qilamiz
    result = result[0].upper() + result[1:] if result else "Nol"
    result += " so'm"

    if tiyin_part:
        result += " " + _three_digits_to_words(tiyin_part) + " tiyin"

    return result


def _make_run(p_elem, text: str, template_rpr, bold: bool):
    """Yangi w:r element yaratadi."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    r_elem = OxmlElement("w:r")

    # Run properties nusxasi
    if template_rpr is not None:
        new_rpr = copy.deepcopy(template_rpr)
    else:
        new_rpr = OxmlElement("w:rPr")

    if bold:
        b_elem = new_rpr.find(qn("w:b"))
        if b_elem is None:
            b_elem = OxmlElement("w:b")
            new_rpr.insert(0, b_elem)
        b_elem.attrib.pop(qn("w:val"), None)
        # bCs ham qo'shish (Kirill/lotin uchun)
        bcs = new_rpr.find(qn("w:bCs"))
        if bcs is None:
            bcs = OxmlElement("w:bCs")
            new_rpr.insert(1, bcs)
    else:
        # bold elementni olib tashlaymiz
        for tag in ("w:b", "w:bCs"):
            el = new_rpr.find(qn(tag))
            if el is not None:
                new_rpr.remove(el)

    r_elem.append(new_rpr)

    t_elem = OxmlElement("w:t")
    t_elem.text = text
    if text != text.strip():
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_elem.append(t_elem)
    p_elem.append(r_elem)


def _is_image_run(run) -> bool:
    """Run rasm (muhr/QR) yoki grafik saqlaydimi?"""
    from docx.oxml.ns import qn

    return run._r.find(qn("w:drawing")) is not None or run._r.find(qn("w:pict")) is not None


def _fill_paragraph(para, placeholders: dict):
    """Paragraph ichidagi placeholder'larni to'ldiradi (rasm/muhr run'larini saqlaydi)."""
    from docx.oxml.ns import qn

    full_text = "".join(run.text for run in para.runs)
    if not any(f"{{{{{k}}}}}" in full_text for k in placeholders):
        return

    has_bold = any(f"{{{{{k}}}}}" in full_text for k in BOLD_PLACEHOLDERS)

    # Bold kerak bo'lmasa — oddiy almashtirish (rasm run'lariga tegmaymiz)
    if not has_bold:
        for key, value in placeholders.items():
            full_text = full_text.replace(f"{{{{{key}}}}}", str(value))
        assigned = False
        for run in para.runs:
            if _is_image_run(run):
                continue
            run.text = full_text if not assigned else ""
            assigned = True
        return

    # Bold kerak: matn run'larini bo'lib qayta yaratamiz, rasm run'larini saqlaymiz
    template_rpr = None
    for run in para.runs:
        if not _is_image_run(run):
            template_rpr = run._r.find(qn("w:rPr"))
            break

    # Faqat matn run'larini o'chirish (rasm/muhr o'z joyida qoladi)
    p_elem = para._p
    for run in list(para.runs):
        if _is_image_run(run):
            continue
        p_elem.remove(run._r)

    # placeholder'larni non-bold bilan almashtir, bold'larni saqla
    for key, value in placeholders.items():
        if key not in BOLD_PLACEHOLDERS:
            full_text = full_text.replace(f"{{{{{key}}}}}", str(value))

    # Bold placeholder'larni pattern bilan bo'l
    bold_pattern = "|".join(
        re.escape(f"{{{{{k}}}}}") for k in BOLD_PLACEHOLDERS if f"{{{{{k}}}}}" in full_text
    )
    parts = re.split(f"({bold_pattern})", full_text) if bold_pattern else [full_text]

    for part in parts:
        if not part:
            continue
        is_bold_part = part in {f"{{{{{k}}}}}" for k in BOLD_PLACEHOLDERS}
        text = placeholders.get(part[2:-2], part) if is_bold_part else part
        _make_run(p_elem, text, template_rpr, bold=is_bold_part)


def _fill_docx(template_path: str, placeholders: dict, post_process=None) -> io.BytesIO:
    """Shablon Word faylni to'ldiradi va BytesIO qaytaradi.

    post_process(docx_doc) — saqlashdan oldin hujjatga qo'shimcha o'zgartirish
    (masalan Kalkulatsiya uchun 1-ilova jadvalini qo'shish) kiritish uchun.
    """
    doc = DocxDocument(template_path)
    for para in doc.paragraphs:
        _fill_paragraph(para, placeholders)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _fill_paragraph(para, placeholders)
    if post_process is not None:
        post_process(doc)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Kalendar reja — bosqich choraklari (boshlanishi/tugashi) va yillik JAMI qatorlari
# ---------------------------------------------------------------------------

_QUARTER_ROMAN = {1: "I", 2: "II", 3: "III", 4: "IV"}

# development_deadline'dan olingan (yil, chorak) → 4 bosqichning
# ((boshlanish_yil, boshlanish_chorak), (tugash_yil, tugash_chorak)) qiymatlari.
# Sheets'da faqat quyidagi 5 ta qiymat uchraydi (foydalanuvchi tasdiqlagan).
_KALENDAR_STAGE_MAP = {
    (2026, 3): [((2026, 3), (2026, 3)), ((2026, 3), (2026, 3)),
                ((2026, 3), (2026, 3)), ((2026, 3), (2026, 3))],
    (2026, 4): [((2026, 3), (2026, 3)), ((2026, 3), (2026, 3)),
                ((2026, 3), (2026, 4)), ((2026, 4), (2026, 4))],
    (2027, 2): [((2026, 3), (2026, 4)), ((2027, 1), (2027, 1)),
                ((2027, 1), (2027, 2)), ((2027, 2), (2027, 2))],
    (2027, 3): [((2026, 3), (2026, 4)), ((2027, 1), (2027, 2)),
                ((2027, 2), (2027, 3)), ((2027, 3), (2027, 3))],
    (2027, 4): [((2026, 3), (2026, 4)), ((2027, 1), (2027, 2)),
                ((2027, 2), (2027, 3)), ((2027, 3), (2027, 4))],
}


def _parse_deadline_quarter(text):
    """'2026-yil IV-chorak' → (2026, 4). Aniqlanmasa None."""
    if not text:
        return None
    m_year = re.search(r"(20\d{2})", text)
    if not m_year:
        return None
    year = int(m_year.group(1))
    up = text.upper()
    for roman, val in (("IV", 4), ("III", 3), ("II", 2), ("I", 1)):
        if re.search(r"(?<![A-Z])" + roman + r"(?![A-Z])", up):
            return (year, val)
    return None


def _fmt_quarter(yq):
    """(2026, 4) → '2026-yil IV-chorak'."""
    year, q = yq
    return f"{year}-yil {_QUARTER_ROMAN[q]}-chorak"


def _compute_kalendar_stages(doc, stage_amounts):
    """development_deadline'dan bosqich choraklari + yillik JAMI guruhlarini hisoblaydi.

    stage_amounts: [I, II, III, IV] bosqich summalari (Decimal).
    Qaytaradi dict yoki None (deadline nomalum bo'lsa):
        {
          "labels": [(start_str, end_str), ...4 ta],
          "stage_years": [yil, ...4 ta],          # har bosqichning boshlanish yili
          "year_totals": [(yil, summa), ...],      # ko'rinish tartibida
          "multi_year": bool,
        }
    """
    parsed = _parse_deadline_quarter(doc.development_deadline)
    if parsed is None or parsed not in _KALENDAR_STAGE_MAP:
        return None
    stages = _KALENDAR_STAGE_MAP[parsed]
    labels = [(_fmt_quarter(s), _fmt_quarter(e)) for (s, e) in stages]
    stage_years = [s[0] for (s, _e) in stages]

    year_totals = {}
    order = []
    for i, year in enumerate(stage_years):
        if year not in year_totals:
            year_totals[year] = Decimal("0.00")
            order.append(year)
        year_totals[year] += stage_amounts[i]

    return {
        "labels": labels,
        "stage_years": stage_years,
        "year_totals": [(y, year_totals[y]) for y in order],
        "multi_year": len(order) > 1,
    }


def _set_cell_text(cell, text):
    """Katak matnini birinchi run formatlashini saqlagan holda almashtiradi."""
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(text)
    for p in cell.paragraphs[1:]:
        for r in p.runs:
            r.text = ""


def _set_jami_row(row, label, amount):
    """JAMI qatorini to'ldiradi: [label, '', '', summa]."""
    cells = row.cells
    _set_cell_text(cells[0], label)
    _set_cell_text(cells[1], "")
    _set_cell_text(cells[2], "")
    _set_cell_text(cells[3], _fmt_money(amount))


# 12-jadval (VHM/sahifa koeffitsienti) — Kalkulatsiya 1-ilovasi uchun.
# Har bir tuple: (Yangi, Qayta, O'zgartirish)  [I/II/III toifa]
_TABLE_12 = {
    "technical_regulation": {"1": (15, 8, 3), "2": (28, 14, 5), "3": (43, 22, 7)},
    "shnq": {"1": (14, 7, 2), "2": (26, 13, 4), "3": (40, 20, 7)},
    "eurocode": {"1": (13, 7, 2), "2": (24, 12, 4), "3": (37, 19, 6)},
    "nizom": {"1": (12, 6, 2), "2": (23, 12, 4), "3": (34, 17, 6)},
    "standard": {"1": (11, 6, 2), "2": (21, 11, 4), "3": (32, 16, 5)},
    "methodical_guide": {"1": (10, 5, 2), "2": (19, 10, 3), "3": (29, 15, 5)},
}

# (label, mos_normativ_turlar, TABLE_12_kaliti)
_JADVAL_12_GROUPS = [
    ("Texnik reglament", ["technical_regulation"], "technical_regulation"),
    ("Shaharsozlik normalari va qoidalari, Milliy qurilish normalari", ["shnq", "mqn"], "shnq"),
    ("Xalqaro yoki xorijiy normalar", ["eurocode"], "eurocode"),
    ("Qurilishda qo'llaniladigan nizom, qoida, yo'riqnoma, tartib va sh.k.", ["nizom"], "nizom"),
    ("Standartlar, Idoraviy qurilish normalari, smeta-resurs normalari", ["standard", "srn", "qr"], "standard"),
    ("Qo'llanma, ma'lumotnoma, boshqa majburiy bo'lmagan hujjatlar", ["methodical_guide"], "methodical_guide"),
]

_LEVEL_ROMAN = {"1": "I", "2": "II", "3": "III"}


def _category_column_index(document_category: str):
    """Hujjat toifasi → jadval ustuni (0=Yangi, 1=Qayta, 2=O'zgartirish)."""
    if document_category == "new":
        return 0
    if document_category in ("rework_harmonization", "rework_modification"):
        return 1
    if document_category == "additional_change":
        return 2
    return None


def _selected_vhm(doc):
    """Hujjat uchun tanlangan VHM/sahifa koeffitsienti (12-jadvaldan) — masalan 14."""
    group = next((g for g in _JADVAL_12_GROUPS if doc.normative_type in g[1]), None)
    col = _category_column_index(doc.document_category)
    level = str(doc.complexity_level)
    if group is None or col is None or level not in _TABLE_12.get(group[2], {}):
        return None
    return _TABLE_12[group[2]][level][col]


# Normativ tur → to'liq nom (Kalkulatsiya sarlavhasi uchun — gap o'rtasida, kichik harf bilan)
NORMATIVE_TYPE_FULL = {
    "shnq": "shaharsozlik normalari va qoidalari",
    "mqn": "milliy qurilish normalari",
    "standard": "milliy standart",
    "srn": "smeta-resurs normalari",
    "qr": "idoraviy qurilish normalari",
    "technical_regulation": "texnik reglament",
    "eurocode": "xalqaro yoki xorijiy normalar",
    "nizom": "nizom, qoida, yo'riqnoma",
    "methodical_guide": "qo'llanma, ma'lumotnoma",
}


def _append_koeffitsient_appendix(docx_doc, doc):
    """Kalkulatsiya hujjati oxiriga 12-jadvalni (1-ilova) qo'shadi.

    Ushbu hujjat uchun tanlangan koeffitsient (turi × toifa × ish turi) ajratib ko'rsatiladi.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    FONT = "Times New Roman"
    LIGHT = "EDEDF7"   # tanlangan guruh — ochiq lavanda
    DARK = "1A227F"    # aynan tanlangan katak — to'q ko'k

    def shade(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        for existing in tcPr.findall(qn("w:shd")):  # eski shading'ni almashtiramiz
            tcPr.remove(existing)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), fill_hex)
        tcPr.append(shd)

    def set_borders(table):
        # "Table Grid" uslubi shablonda bo'lmasligi mumkin — chegarani XML orqali qo'yamiz
        tblPr = table._tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "999999")
            borders.append(el)
        tblPr.append(borders)

    def valign_center(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        for existing in tcPr.findall(qn("w:vAlign")):
            tcPr.remove(existing)
        v = OxmlElement("w:vAlign")
        v.set(qn("w:val"), "center")
        tcPr.append(v)

    def set_row_height(row, twips, rule="atLeast"):
        trPr = row._tr.get_or_add_trPr()
        h = OxmlElement("w:trHeight")
        h.set(qn("w:val"), str(twips))
        h.set(qn("w:hRule"), rule)
        trPr.append(h)

    def set_cell(cell, text, *, bold=False, color=None, align="left", size=8):
        cell.text = ""
        valign_center(cell)  # matn katakda vertikal markazda
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = FONT
        if color:
            run.font.color.rgb = RGBColor.from_string(color)

    # Tanlangan pozitsiya
    sel_group = next(
        (i for i, (_, types, _) in enumerate(_JADVAL_12_GROUPS) if doc.normative_type in types),
        None,
    )
    sel_level = str(doc.complexity_level)
    sel_col = _category_column_index(doc.document_category)  # 0/1/2 yoki None

    # Jadval "Nb – ... qiymat (N);" qatoridan keyin joylashadi (o'sha qiymat qayerdan
    # kelganini ko'rsatuvchi javob). Ankor sifatida Nb paragrafini topamiz.
    anchor = None
    for p in docx_doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Nb ") and ("qiymat" in t or "me'yorlariga" in t):
            anchor = p._p
            break

    created = []  # yangi qo'shilgan elementlar (keyin Nb tagiga ko'chiramiz)

    # --- Jadval ---
    headers = ["Hujjat turi", "Murakkablik darajasi", "Yangi", "Qayta", "O'zgartirish"]
    table = docx_doc.add_table(rows=1, cols=len(headers))
    created.append(table._tbl)
    set_borders(table)
    for ci, htext in enumerate(headers):
        set_cell(table.rows[0].cells[ci], htext, bold=True, align="center", size=9)
        shade(table.rows[0].cells[ci], "D9D9E8")
    set_row_height(table.rows[0], 340)

    # Faqat ushbu hujjatga tegishli guruhni ko'rsatamiz (tur aniqlansa);
    # aks holda barcha guruhlar (zaxira).
    if sel_group is not None:
        groups_to_show = [(sel_group, _JADVAL_12_GROUPS[sel_group])]
    else:
        groups_to_show = list(enumerate(_JADVAL_12_GROUPS))

    for gi, (label, _types, tkey) in groups_to_show:
        first_row_cells = None
        group_rows = []
        for level in ("1", "2", "3"):
            row = table.add_row()
            set_row_height(row, 340)  # barcha qatorlar teng va past balandlikda (~0.6 sm)
            group_rows.append(row)
            cells = row.cells
            if first_row_cells is None:
                first_row_cells = cells
            values = _TABLE_12[tkey][level]  # (Yangi, Qayta, O'zgartirish)
            # ustun 0 — nom (keyin vertikal birlashtiramiz)
            set_cell(cells[0], label if level == "1" else "", size=8)
            set_cell(cells[1], _LEVEL_ROMAN[level], align="center", size=8)
            for k in range(3):
                set_cell(cells[2 + k], str(values[k]), align="center", size=8)

            is_sel_group = gi == sel_group
            if is_sel_group:
                for c in cells:
                    shade(c, LIGHT)
                # aynan tanlangan katak
                if sel_col is not None and level == sel_level:
                    target = cells[2 + sel_col]
                    shade(target, DARK)
                    set_cell(target, str(values[sel_col]), bold=True, color="FFFFFF", align="center", size=9)

        # nom ustunini (0) 3 qatorga birlashtirish — matn katak markazida (vertikal + gorizontal)
        merged = first_row_cells[0]
        for r2 in group_rows[1:]:
            merged = merged.merge(r2.cells[0])
        merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        valign_center(merged)

    # --- Tanlov izohi ---
    if sel_group is not None and sel_col is not None:
        cat_name = {0: "Yangi", 1: "Qayta", 2: "O'zgartirish"}[sel_col]
        note = docx_doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = note.add_run(
            f"Ushbu normativ hujjat uchun tanlangan: "
            f"Murakkablik darajasi {_LEVEL_ROMAN[sel_level]}, {cat_name} - "
            f"{_TABLE_12[_JADVAL_12_GROUPS[sel_group][2]][sel_level][sel_col]}"
        )
        r.italic = True
        r.font.name = FONT
        r.font.size = Pt(9)
        created.append(note._p)

    # --- Yangi qo'shilgan elementlarni Nb qatoridan keyinga ko'chiramiz ---
    if anchor is not None:
        ref = anchor
        for el in created:
            ref.addnext(el)  # lxml elementni ko'chiradi (nusxa emas)
            ref = el


class DocumentContractAPIView(APIView):
    """Shartnoma shablonini to'ldirib base64 .docx qaytaradi."""

    authentication_classes = []
    permission_classes = []

    TEMPLATE_NAME = "shartnoma.docx"

    def get(self, request, pk):
        doc = get_object_or_404(DocumentCalculation, pk=pk)

        template_path = os.path.join(
            settings.BASE_DIR, "contract_templates", self.TEMPLATE_NAME
        )
        if not os.path.exists(template_path):
            return Response(
                {"error": f"Shablon fayl topilmadi: {self.TEMPLATE_NAME}"},
                status=status.HTTP_404_NOT_FOUND,
            )

        org = OrganizationSettings.get_instance()

        # Formuladan to'g'ri yakuniy summani hisoblash (DB'dagi eski qiymat emas)
        # apply_normative_coefficients() → NormativeCoefficient jadvalidan VHM oladi
        # recalculate_final_total_amount() → VHM × pages × MROT × 2.1 × 1.12 [× 1.4]
        doc.apply_normative_coefficients()
        doc.recalculate_final_total_amount()
        # doc.final_total_amount endi to'g'ri hisoblangan qiymat (DB'ga saqlanmaydi)

        # 2.1-band uchun yillar bo'yicha taqsimot (Google Sheets'dan keladigan qiymatlar):
        #   2026-yilga = planned_amount (2026-yilga rejalashtirilgan)
        #   2027-yil uchun = umumiy - 2026 - 01.01.2026 holatiga bajarilgan (completed_amount)
        # 01.01.2026 holatidagi summa 2025-yil uchun bajarilgan hisoblanadi.
        # Sheets'da summalar ming so'mda saqlanadi — umumiy (formula) so'mda bo'lgani uchun ×1000.
        SHEET_SCALE = Decimal("1000")
        amount_2026 = (doc.planned_amount or Decimal("0.00")) * SHEET_SCALE
        completed_2025 = (doc.completed_amount or Decimal("0.00")) * SHEET_SCALE
        amount_2027 = doc.final_total_amount - amount_2026 - completed_2025
        if amount_2027 < 0:
            amount_2027 = Decimal("0.00")

        # Kalkulatsiya uchun: MHb (bazaviy narx) va MHi (ilmiy-tadqiqot narxi)
        #   final = base × 1.4 (agar tadqiqot bo'lsa) → base = final / 1.4, MHi = final - base
        #   tadqiqot bo'lmasa → base = final, MHi = 0
        research_factor = Decimal("1.4") if doc.is_research_required else Decimal("1")
        mhb_amount = (doc.final_total_amount / research_factor).quantize(
            Decimal("0.01"), rounding=ROUND_HALF_UP
        )
        mhi_amount = doc.final_total_amount - mhb_amount
        if mhi_amount < 0:
            mhi_amount = Decimal("0.00")

        # Kalendar reja bosqich summalari: I=30%, II=30%, III=30%, IV=10% (oxirgisi qoldiq)
        _kr_total = doc.final_total_amount
        kr_stage1 = (_kr_total * Decimal("0.30")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
        kr_stage2 = (_kr_total * Decimal("0.30")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
        kr_stage3 = (_kr_total * Decimal("0.30")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
        kr_stage4 = _kr_total - kr_stage1 - kr_stage2 - kr_stage3  # ~10%, yig'indi umumiyga teng

        # Kalendar reja bosqich choraklari — development_deadline'dan avtomatik hisoblanadi
        kr_info = _compute_kalendar_stages(
            doc, [kr_stage1, kr_stage2, kr_stage3, kr_stage4]
        )

        placeholders = {
            # Hujjat ma'lumotlari
            "shnq_name": doc.name,
            "doc_id": str(doc.id),
            "normative_type": doc.normative_type,
            "document_category": doc.document_category,
            "total_pages": str(doc.total_pages),
            "complexity_level": str(doc.complexity_level),
            # TZ (Texnik topshiriq) uchun — o'qiladigan ko'rinishdagi qiymatlar
            "complexity_label": doc.get_complexity_level_display(),
            "complexity_roman": {"1": "I", "2": "II", "3": "III"}.get(str(doc.complexity_level), str(doc.complexity_level)),
            "document_category_label": doc.get_document_category_display(),
            "sources_count": str(doc.sources_count),
            "research_status": "Ha" if doc.is_research_required else "Yo'q",
            # Nb izohi uchun — 12-jadvaldan tanlangan VHM/sahifa koeffitsienti (masalan 14)
            "vhm_value": str(_selected_vhm(doc) if _selected_vhm(doc) is not None else ""),
            # Kalkulatsiya: normativ tur to'liq nomi + MHb/MHi summalari
            "normative_type_full": NORMATIVE_TYPE_FULL.get(doc.normative_type, "normativ hujjat"),
            "mhb_amount": _fmt_money(mhb_amount),
            "mhi_amount": _fmt_money(mhi_amount),
            "executor_organization": doc.executor_organization or "",
            "development_deadline": doc.development_deadline or "",
            "shartnoma_number": doc.contract_number or f"{doc.id}/26",
            "current_year_percent": str(doc.current_year_percent),
            "next_year_percent": str(max(Decimal("0"), Decimal("100") - doc.current_year_percent)),
            "current_year_amount": _fmt_money(
                doc.final_total_amount * doc.current_year_percent / Decimal("100")
            ),
            "next_year_amount": _fmt_money(
                doc.final_total_amount * max(Decimal("0"), Decimal("100") - doc.current_year_percent) / Decimal("100")
            ),
            "notes": doc.notes or "",
            "final_total_amount": _fmt_money(doc.final_total_amount),
            "final_total_amount_words": f"({number_to_uz_words(doc.final_total_amount)})",
            # 2.1-band: yillar bo'yicha taqsimot
            "amount_2026": _fmt_money(amount_2026),
            "amount_2027": _fmt_money(amount_2027),
            "created_at": doc.created_at.strftime("%d.%m.%Y") if doc.created_at else "",
            # Tashkilot sozlamalari
            "institute_director": org.institute_director,
            "deputy_minister": org.deputy_minister,
            "economics_head": org.economics_head,
            # Kalendar reja bosqichlari — summalar umumiy narxdan 30/30/30/10 taqsimlanadi
            "I_start": doc.stage1_start,
            "I_end": doc.stage1_end,
            "I_summa": _fmt_money(kr_stage1),
            "II_start": doc.stage2_start,
            "II_end": doc.stage2_end,
            "II_summa": _fmt_money(kr_stage2),
            "III_start": doc.stage3_start,
            "III_end": doc.stage3_end,
            "III_summa": _fmt_money(kr_stage3),
            "IV_start": doc.stage4_start,
            "IV_end": doc.stage4_end,
            "IV_summa": _fmt_money(kr_stage4),
        }

        # development_deadline aniqlansa — bosqich choraklarini hisoblangan qiymat bilan almashtiramiz
        if kr_info is not None:
            for key, (start, end) in zip(("I", "II", "III", "IV"), kr_info["labels"]):
                placeholders[f"{key}_start"] = start
                placeholders[f"{key}_end"] = end

        buf = _fill_docx(
            template_path,
            placeholders,
            post_process=lambda docx_doc: self.post_process_docx(docx_doc, doc),
        )
        encoded = base64.b64encode(buf.read()).decode("utf-8")

        return Response(
            {
                "filename": f"shartnoma_{doc.id}.docx",
                "doc_name": doc.name,
                "data": encoded,
            },
            status=status.HTTP_200_OK,
        )

    def post_process_docx(self, docx_doc, doc):
        """Sub-view'lar hujjatga qo'shimcha element qo'shishi uchun.

        Shartnoma uchun: 2027 summasiga qarab "Shundan..." jumlasi va 1.3-band muddati.
        Kalendar reja uchun: yillar bo'yicha JAMI qatorlarini joylashtirish.
        """
        if self.TEMPLATE_NAME == "kalendar_reja.docx":
            self._post_process_kalendar(docx_doc, doc)
            return None
        if self.TEMPLATE_NAME != "shartnoma.docx":
            return None

        scale = Decimal("1000")
        amount_2026 = (doc.planned_amount or Decimal("0")) * scale
        completed_2025 = (doc.completed_amount or Decimal("0")) * scale
        amount_2027 = doc.final_total_amount - amount_2026 - completed_2025

        if amount_2027 <= 0:
            # 2027 uchun summa yo'q → "Shundan 2026-yilga ..." jumlasini olib tashlaymiz
            for p in docx_doc.paragraphs:
                if p.text.strip().startswith("Shundan 2026-yilga"):
                    p._p.getparent().remove(p._p)
                    break
        else:
            # 2027 uchun summa bor → 1.3-band muddati 2026-yil → 2027-yil 31-dekabrgacha
            for p in docx_doc.paragraphs:
                if "31-dekabrgacha" not in p.text:
                    continue
                runs = p.runs
                anchor_i = next(
                    (i for i, r in enumerate(runs) if "-yil" in r.text and "dekabr" in r.text),
                    None,
                )
                if anchor_i and anchor_i > 0:
                    prev = runs[anchor_i - 1]
                    if prev.text.endswith("6"):  # yil oxirgi raqami: ...2026 → ...2027
                        prev.text = prev.text[:-1] + "7"
                break
        return None

    def _post_process_kalendar(self, docx_doc, doc):
        """Kalendar reja jadvalidagi JAMI qatorini yillar bo'yicha joylashtiradi.

        Bir yil (2026'da tugasa): mavjud JAMI qatori "JAMI: 2026-yil uchun".
        Ikki yil (2027'ga o'tsa): I-bosqichdan keyin "JAMI: 2026-yil uchun",
        oxirida "JAMI: 2027-yil uchun". Summalar bosqichlar yig'indisi.
        """
        total = doc.final_total_amount
        s1 = (total * Decimal("0.30")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
        s2 = (total * Decimal("0.30")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
        s3 = (total * Decimal("0.30")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
        s4 = total - s1 - s2 - s3
        info = _compute_kalendar_stages(doc, [s1, s2, s3, s4])
        if info is None:
            return

        # JAMI qatori bo'lgan jadvalni topamiz
        table = None
        jami_idx = None
        for t in docx_doc.tables:
            for ri, row in enumerate(t.rows):
                if row.cells[0].text.strip().upper().startswith("JAMI"):
                    table, jami_idx = t, ri
                    break
            if table is not None:
                break
        if table is None:
            return

        year_totals = info["year_totals"]
        stage_years = info["stage_years"]

        # Mavjud (oxirgi) JAMI qatori → oxirgi yil guruhi
        last_year, last_amount = year_totals[-1]
        _set_jami_row(table.rows[jami_idx], f"JAMI: {last_year}-yil uchun", last_amount)

        if not info["multi_year"]:
            return

        # Bosqich qatorlari: jami_idx-4 .. jami_idx-1 (I, II, III, IV)
        stage_row_trs = [table.rows[jami_idx - 4 + i]._tr for i in range(4)]
        jami_tr = table.rows[jami_idx]._tr  # nusxa uchun namuna (formatlash saqlanadi)

        # Oxirgidan tashqari har bir yil guruhi uchun: guruhning oxirgi bosqichidan
        # keyin JAMI qatorini kiritamiz
        for year, amount in year_totals[:-1]:
            last_stage_i = max(i for i, y in enumerate(stage_years) if y == year)
            new_tr = copy.deepcopy(jami_tr)
            stage_row_trs[last_stage_i].addnext(new_tr)
            new_row = _TableRow(new_tr, table)
            _set_jami_row(new_row, f"JAMI: {year}-yil uchun", amount)


class DocumentKalendarRejaAPIView(DocumentContractAPIView):
    """Kalendar reja shablonini to'ldirib base64 .docx qaytaradi."""

    TEMPLATE_NAME = "kalendar_reja.docx"

    def get(self, request, pk):
        response = super().get(request, pk)
        if response.status_code == 200:
            response.data["filename"] = f"kalendar_reja_{pk}.docx"
        return response


class DocumentTexnikTopshiriqAPIView(DocumentContractAPIView):
    """Texnik topshiriq (TZ) shablonini to'ldirib base64 .docx qaytaradi."""

    TEMPLATE_NAME = "TZ.docx"

    def get(self, request, pk):
        response = super().get(request, pk)
        if response.status_code == 200:
            response.data["filename"] = f"TZ_{pk}.docx"
        return response


class DocumentBayonnomaAPIView(DocumentContractAPIView):
    """Bayonnoma (kelishuv qiymati to'g'risida) shablonini to'ldirib base64 .docx qaytaradi."""

    TEMPLATE_NAME = "bayonnoma.docx"

    def get(self, request, pk):
        response = super().get(request, pk)
        if response.status_code == 200:
            response.data["filename"] = f"bayonnoma_{pk}.docx"
        return response


class DocumentKalkulatsiyaAPIView(DocumentContractAPIView):
    """Kalkulatsiya shablonini to'ldirib base64 .docx qaytaradi."""

    TEMPLATE_NAME = "Kalkul.docx"

    def get(self, request, pk):
        response = super().get(request, pk)
        if response.status_code == 200:
            response.data["filename"] = f"kalkulatsiya_{pk}.docx"
        return response

    def post_process_docx(self, docx_doc, doc):
        # Kalkulatsiya oxiriga 12-jadvalni (1-ilova) tanlangan koeffitsient ajratilgan holda qo'shamiz
        _append_koeffitsient_appendix(docx_doc, doc)


class DocumentKalkulatsiyaBulkAPIView(APIView):
    """Barcha hujjatlarning Kalkulatsiya .docx fayllarini bitta ZIP qilib qaytaradi."""

    authentication_classes = []
    permission_classes = []

    def get(self, request):
        view = DocumentKalkulatsiyaAPIView()
        used_names = set()

        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for doc in DocumentCalculation.objects.all().order_by("id"):
                try:
                    resp = view.get(request, doc.pk)
                except Exception:
                    continue
                if getattr(resp, "status_code", None) != 200:
                    continue
                docx_bytes = base64.b64decode(resp.data["data"])
                safe = _re_module.sub(r"[^\w\s.\-]", "", doc.name or "").strip()[:80] or "hujjat"
                fname = f"{doc.id}_{safe}.docx"
                # nomlar takrorlanmasin
                base, ext = os.path.splitext(fname)
                n = 1
                while fname in used_names:
                    fname = f"{base}_{n}{ext}"
                    n += 1
                used_names.add(fname)
                zf.writestr(fname, docx_bytes)

        zip_buf.seek(0)
        response = HttpResponse(zip_buf.getvalue(), content_type="application/zip")
        response["Content-Disposition"] = 'attachment; filename="kalkulatsiyalar.zip"'
        return response


class SyncFromSheetsAPIView(APIView):
    """Google Sheets dan bazani qo'lda yangilash."""

    authentication_classes = []
    permission_classes = []

    def post(self, request):
        from django.core.management import call_command
        import io as _io

        out = _io.StringIO()
        try:
            call_command("sync_from_sheets", stdout=out, stderr=out)
        except Exception as exc:
            return Response(
                {"detail": f"Sync xatosi: {exc}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

        output = out.getvalue()
        # "Yaratildi: N ta" satridan natijani olish
        created = 0
        for line in output.splitlines():
            if "Yaratildi:" in line:
                parts = line.split("Yaratildi:")
                if len(parts) > 1:
                    try:
                        created = int(parts[1].strip().split()[0])
                    except (ValueError, IndexError):
                        pass

        return Response(
            {"detail": "Muvaffaqiyatli yangilandi.", "created": created, "log": output},
            status=status.HTTP_200_OK,
        )


class OrganizationSettingsAPIView(APIView):
    """GET — sozlamalarni olish, PUT — yangilash."""

    authentication_classes = []
    permission_classes = []

    def get(self, request):
        obj = OrganizationSettings.get_instance()
        return Response(OrganizationSettingsSerializer(obj).data)

    def put(self, request):
        obj = OrganizationSettings.get_instance()
        serializer = OrganizationSettingsSerializer(obj, data=request.data, partial=True)
        serializer.is_valid(raise_exception=True)
        serializer.save()
        return Response(serializer.data)
