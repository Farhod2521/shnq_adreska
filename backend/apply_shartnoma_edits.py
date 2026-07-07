"""
shartnoma.docx ga ikkita o'zgarish kiritadi (Word yopiq bo'lishi kerak):
  1) Direktor imzosi ({{institute_director}}) oldiga pechat.png muhrini qo'yadi.
  2) "deb belgilandi." paragrafidan keyin yillar bo'yicha taqsimot jumlasini qo'shadi:
     "Shundan 2026-yilga {{amount_2026}} so'm, 2027-yil uchun esa {{amount_2027}} so'm."

Ishlatish:  python apply_shartnoma_edits.py
"""

import copy
import os

import docx
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.text.paragraph import Paragraph

DIR = os.path.join(os.path.dirname(__file__), "contract_templates")
PATH = os.path.join(DIR, "shartnoma.docx")
IMG = os.path.join(DIR, "pechat.png")
SENTENCE = "\tShundan 2026-yilga {{amount_2026}} so'm, 2027-yil uchun esa {{amount_2027}} so'm."


def insert_seal(para):
    runs = para.runs
    u = next((i for i, r in enumerate(runs) if "____" in r.text), None)
    if u is None:
        return False
    anchor = runs[u + 1] if u + 1 < len(runs) else runs[u]
    new_run = para.add_run()
    new_run.add_picture(IMG, width=Cm(2.2))
    anchor._r.addprevious(new_run._r)
    return True


def main():
    # yozib bo'lishini tekshirish
    try:
        fh = open(PATH, "ab")
        fh.close()
    except PermissionError:
        print("shartnoma.docx BAND (Word ochiq). Word'ni yopib qayta ishga tushiring.")
        return

    d = docx.Document(PATH)

    # 1) Muhr — faqat imzo qatorida ('____' + {{institute_director}})
    seals = 0
    for p in d.paragraphs:
        if "{{institute_director}}" in p.text and "____" in p.text:
            if insert_seal(p):
                seals += 1
    for t in d.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    if "{{institute_director}}" in p.text and "____" in p.text:
                        if insert_seal(p):
                            seals += 1

    # 2) "deb belgilandi." dan keyin taqsimot jumlasi
    added = False
    for p in d.paragraphs:
        if p.text.strip() == "deb belgilandi.":
            new_el = copy.deepcopy(p._p)  # formatlashni nusxalaymiz
            p._p.addnext(new_el)
            new_para = Paragraph(new_el, p._parent)
            runs = new_para.runs
            if runs:
                runs[0].text = SENTENCE
                for r in runs[1:]:
                    r.text = ""
            else:
                new_para.add_run(SENTENCE)
            added = True
            break

    d.save(PATH)
    print(f"Muhr qo'yildi: {seals} ta imzo qatoriga")
    print(f"Taqsimot jumlasi qo'shildi: {added}")


if __name__ == "__main__":
    main()
